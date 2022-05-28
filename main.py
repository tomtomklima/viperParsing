import openpyxl
import docx
import docx2pdf
from os import listdir
from os.path import isfile, join

class Defect:
    def __init__(
            self,
            id_number,
            name,
            description,
            description_note,
            weight,
            place_a,
            place_b,
            place_c,
            severity,
            repair,
            photo_id,
    ):
        self.id_number = str(id_number.value)
        self.name = str(name.value)
        self.description = str(description.value)
        self.description_note = str(description_note.value)
        self.weight = str(weight.value)
        self.place_a = str(place_a.value)
        self.place_b = str(place_b.value)
        self.place_c = str(place_c.value)
        self.severity = str(severity.value)
        self.repair_a = "☒" if str(repair.value) == "O - oprava" else "☐"
        self.repair_b = "☒" if str(repair.value) == "N - výměna" else "☐"
        self.repair_c = "☒" if str(repair.value) == "N - sleva" else "☐"
        self.photo_id = str(photo_id.value)


def main(path):
    sheet = openpyxl.load_workbook(path).active

    photos_path = "files/photos/"
    photos = [f for f in listdir(photos_path) if isfile(join(photos_path, f))]

    
    defects = []
    for row_id in range(4, 9999):
        if sheet.cell(row_id, 1).value is None:
            break
        
        defect = Defect(
            sheet.cell(row_id, 14),
            sheet.cell(row_id, 9),
            sheet.cell(row_id, 10),
            sheet.cell(row_id, 11),
            sheet.cell(row_id, 7),
            sheet.cell(row_id, 4),
            sheet.cell(row_id, 5),
            sheet.cell(row_id, 6),
            sheet.cell(row_id, 12),
            sheet.cell(row_id, 13),
            sheet.cell(row_id, 3),
        )
        defects.append(defect)
    
    for defect in defects:
        document = docx.Document("files/template.docx")
        
        mapper = {
            "%0_reclamation_number%": "id_number",
            "%1_name%": "name",
            "%2_place_a%": "place_a",
            "%3_place_b%": "place_b",
            "%4_place_c%": "place_c",
            "%5_description%": "description",
            "%6_description_note%": "description_note",
            "%7_severity%": "severity",
            "%8_repair_a%": "repair_a",
            "%9_repair_b%": "repair_b",
            "%10_repair_c%": "repair_c",
        }
        for keyword, replacing_index in mapper.items():
            for paragraph in document.paragraphs:
                if paragraph.text.find(keyword) != -1:
                    # added run for keeping style
                    for run in paragraph.runs:
                        inline_position = run.text.find(keyword)
                        if inline_position != -1:
                            new_value = str(getattr(defect, replacing_index)) or ""
                            run.text = run.text.replace(keyword, new_value)
                            print(keyword + " -> " + new_value)

        for photo_name in photos:
            if photo_name.startswith(defect.photo_id):
                document.add_paragraph(photo_name + "\r")
                document.add_picture(photos_path + photo_name, width=4.8*914400)
                document.add_paragraph("\r")
                print("added photo " + photo_name)

        new_name = "ELI II Zápis o reklamaci č. " + defect.id_number
        document.save("files/output/" + new_name + ".docx")
        print("Saved document named: " + new_name)
        docx2pdf.convert("files/output/" + new_name + ".docx", "files/output/" + new_name + ".pdf")


if __name__ == '__main__':
    main("files/export.xlsx")
